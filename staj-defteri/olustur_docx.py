# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path(__file__).parent
IMG = BASE / "gorseller"
OUT = BASE / "Staj_Defteri_10_Gun.docx"

DAYS = [
    {
        "title": "1. Gün — Oryantasyon ve Proje Kurulumu",
        "image": "gun-01-oryantasyon.png",
        "konu": "Belediye tanıtımı, proje planlaması ve site kurulumu",
        "ana_hatlar": "Ekip tanıtımı · Proje hedefleri · Site altyapısı · İlk test",
        "text": (
            "Stajın ilk gününde Bilgi İşlem Müdürlüğü'ne gittim, ekiple tanıştım. "
            "Görevim vatandaşların internetten başvuru yapabileceği, başvurularını "
            "takip edebileceği ve harita üzerinden şikayet iletebileceği bir "
            "e-Belediye sitesi geliştirmekti. Belediyenin mevcut sitesini inceledim; "
            "müdürlükleri, duyuruları ve hizmetleri not aldım. Eski deneme projesi "
            "yerine daha düzenli bir web sitesi altyapısı kurdum. Belediye renklerine "
            "uygun görünüm ayarlarını yaptım, ortak buton ve form stillerini "
            "hazırladım. Site başlığını ve mobil uyumlu yapıyı ayarladım. "
            "Bilgisayarımda çalıştırıp ana sayfanın düzgün açıldığını kontrol ettim."
        ),
    },
    {
        "title": "2. Gün — Üst Menü, Alt Bilgi ve Ana Sayfa",
        "image": "gun-04-anasayfa.png",
        "konu": "Site iskeleti, menü yapısı ve ana sayfa düzeni",
        "ana_hatlar": "Üst menü · Alt bilgi · Mobil menü · Hizmet kartları · Müdürlükler",
        "text": (
            "Sitenin her sayfasında görünecek üst menü ve alt bilgi bölümünü "
            "hazırladım. Kullanıcı hangi sayfadaysa menüde o bağlantıyı "
            "belirgin gösterdim. Telefonda yer kazanmak için üç çizgili mobil "
            "menü ekledim; tıklayınca sağdan açılıyor. Menü linklerini tek "
            "bir listeden çektim, yeni sayfa eklemek kolay olsun diye böyle "
            "yaptım. Ana sayfada tanıtım alanı, hizmet kartları ve canlı "
            "istatistik bölümünü oluşturdum. Site içi arama kutusunu ekledim. "
            "Müdürlükler sayfasında 13 müdürlüğü listeledim, tıklayınca "
            "detay sayfasına gidiyor."
        ),
    },
    {
        "title": "3. Gün — Veritabanı ve Online Başvuru",
        "image": "gun-05-veritabani.png",
        "konu": "Veri kaydı, başvuru formu ve dosya ekleme",
        "ana_hatlar": "Bulut veritabanı · Başvuru formu · Alan kontrolü · Dosya yükleme",
        "text": (
            "Başvuruların kaybolmaması için internet üzerinde çalışan bir "
            "veritabanı kurdum. Başvuru, duyuru, randevu ve doğrulama kodu "
            "kayıtları için ayrı tablolar oluşturdum. Harita şikayetlerinde "
            "konum ve adres bilgisi de saklansın diye ek alanlar ekledim. "
            "Veritabanına kayıt ekleme, okuma ve güncelleme işlemlerini "
            "yazdım. Başvuru sayfasına TC, ad soyad, telefon, departman, "
            "konu ve açıklama alanlarını koydum. Form gönderilince bilgileri "
            "kontrol ediyorum; eksik veya hatalıysa uyarı veriyorum. Dosya "
            "eklenmişse buluta yükleyip bağlantısını kaydediyorum. Başvuru "
            "başarılı olunca otomatik numara üretip vatandaşa gösteriyorum."
        ),
    },
    {
        "title": "4. Gün — Başvuru Sorgulama ve Duyurular",
        "image": "gun-07-sorgulama.png",
        "konu": "Başvuru takibi, sayaçlar ve duyuru sayfaları",
        "ana_hatlar": "Sorgulama · Durum etiketleri · Ana sayfa sayaçları · Arama · Duyurular",
        "text": (
            "Vatandaşlar başvuru numarasıyla takip edebilsin diye sorgulama "
            "sayfasını geliştirdim. Numara girilince kaydı bulup durum, "
            "departman, konu, tarih ve varsa konum bilgisini gösteriyorum. "
            "Durumları renkli etiketlerle ayırdım: İncelemede turuncu, "
            "Devam Ediyor kırmızı, Çözüldü yeşil. Ana sayfadaki sayaçlar "
            "veritabanındaki gerçek sayıları gösteriyor, yüklenirken "
            "animasyonlu artıyor. Duyuru listeleme ve detay sayfalarını "
            "veritabanına bağladım. Klavyeden kısayolla açılan arama "
            "penceresini ekledim; yazdıkça site içinde arama yapılıyor."
        ),
    },
    {
        "title": "5. Gün — Harita Tabanlı Şikayet Sistemi",
        "image": "gun-09-harita.png",
        "konu": "Harita ekranı ve konum seçimi",
        "ana_hatlar": "Harita · Şikayet türleri · Dörtyol sınırı · Başvuru yönlendirme",
        "text": (
            "Su kesintisi, elektrik arızası ve bozuk yol şikayetleri için "
            "harita özelliğini siteye ekledim. Harita ekranını hazırladım; "
            "solda konum seçimi ve şikayet türü, sağda açık kaynaklı harita "
            "görünüyor. Üç şikayet tipini tanımladım. Vatandaş haritada "
            "Dörtyol dışına tıklarsa uyarı veriyorum, sadece ilçe içindeki "
            "noktalar kabul ediliyor. Konum ve şikayet türü seçilince "
            "bilgileri geçici olarak saklayıp ayrı bir başvuru sayfasına "
            "yönlendirdim. Böylece harita ekranı sade kalıyor, form ayrı "
            "sayfada dolduruluyor."
        ),
    },
    {
        "title": "6. Gün — Otomatik Adres Bulma",
        "image": "gun-10-adres.png",
        "konu": "Haritadan seçilen noktanın adrese çevrilmesi",
        "ana_hatlar": "Koordinat · Adres doldurma · Yedek servis · İletişim formu",
        "text": (
            "Haritaya tıklayınca seçilen noktanın adresinin otomatik yazılması "
            "için bir servis bağladım. Enlem ve boylam bilgisini dış bir "
            "adres servisine gönderiyorum, dönen cevaptan mahalle ve cadde "
            "bilgisini birleştirip adres alanını dolduruyorum. Su ve elektrik "
            "şikayetlerinde daha ayrıntılı, bozuk yol şikayetlerinde daha "
            "kısa adres gösteriyorum. Bir servis cevap vermezse yedek "
            "servise geçiyorum. Adres alınırken ekranda yükleniyor yazısı "
            "çıkıyor, vatandaş isterse adresi elle de düzenleyebiliyor. "
            "İletişim formunu ve e-posta abonelik bölümünü de bu gün tamamladım."
        ),
    },
    {
        "title": "7. Gün — Yönetici Paneli ve Giriş Sistemi",
        "image": "gun-11-admin.png",
        "konu": "Personel girişi ve başvuru yönetimi",
        "ana_hatlar": "Şifreli giriş · Başvuru listesi · Durum güncelleme · Excel aktarımı",
        "text": (
            "Belediye personeli için şifre korumalı yönetici panelini yaptım. "
            "Doğru kullanıcı adı ve şifre girilince oturum açılıyor, yaklaşık "
            "24 saat geçerli kalıyor. Panele her girişte yetki kontrolü "
            "yapıyorum. Başvuruları listeledim, departmana göre "
            "filtreleyebiliyorum. Durum, not ve atanan personel bilgisini "
            "güncelleyebiliyorum. Listeyi Excel'e aktarma özelliğini ekledim; "
            "Türkçe karakterler bozulmasın diye uygun dosya formatını "
            "kullandım. Yönetici girişini ayrı bir sayfaya taşıdım."
        ),
    },
    {
        "title": "8. Gün — Özet Ekran, Duyuru Yönetimi ve Ek Hizmetler",
        "image": "gun-12-admin-icerik.png",
        "konu": "Yönetici özeti, harita görünümü, randevu ve yardım asistanı",
        "ana_hatlar": "Özet ekran · Harita pinleri · Duyuru yönetimi · Doğrulama kodu · Randevu",
        "text": (
            "Yönetici paneline özet ekran ekledim; toplam başvuru, durum "
            "dağılımı ve şikayet türlerine göre sayıları gösteriyorum. "
            "Konum bilgisi olan başvuruları harita üzerinde işaret olarak "
            "gösterdim, renk şikayet türüne göre değişiyor. Personelin "
            "duyuru ekleyip düzenleyebileceği bir yönetim bölümü yaptım. "
            "Başvuru geçmişine bakmak için e-posta doğrulama kodu sistemi "
            "kurdum; 6 haneli kod gönderiliyor, doğrulandıktan sonra "
            "geçmiş başvurular görünüyor. Randevu alma modülünü ve "
            "sitedeki yardım asistanını da ekledim. KVKK ve gizlilik "
            "sayfalarını yazdım."
        ),
    },
    {
        "title": "9. Gün — Karanlık Mod, Bildirim ve Vatandaş Kayıt",
        "image": "gun-03-menu.png",
        "konu": "Tema seçimi, bildirimler ve üyelik sistemi",
        "ana_hatlar": "Karanlık mod · Bildirim zili · Kayıt ol · Giriş yap",
        "text": (
            "Kullanıcıların istediğinde karanlık moda geçebilmesi için "
            "tema sistemini kurdum. Tercih tarayıcıda saklanıyor, site "
            "varsayılan olarak açık modda açılıyor. Üst menüye ay ve güneş "
            "ikonlu geçiş düğmesi ekledim. Kart, form ve menü renklerini "
            "karanlık moda uygun hale getirdim. Başvuru durumu değişince "
            "vatandaşa site içi bildirim düşsün diye bildirim sistemini "
            "kurdum; üst menüde zil ikonu ve okunmamış sayısı görünüyor. "
            "Vatandaşların hesap açıp giriş yapabilmesi için kayıt ve "
            "giriş sayfalarını hazırladım. Şifreler güvenli şekilde "
            "saklanıyor. İşlem sonrası kısa bilgi mesajları da gösteriliyor."
        ),
    },
    {
        "title": "10. Gün — Test, Yayın ve Proje Sunumu",
        "image": "gun-14-mobil-yayin.png",
        "konu": "Mobil uyum, test verisi, internete yükleme ve sunum",
        "ana_hatlar": "Mobil görünüm · Test verisi · Tüm özellikleri deneme · Canlı sunum",
        "text": (
            "Mobil görünümü iyileştirdim; kaydırmalı yan menü ve kart "
            "tabanlı meclis listesi ekledim. Panel boş görünmesin diye "
            "yaklaşık 80 örnek başvuru ve 20 randevu üreten test verisi "
            "hazırladım. Veritabanı bağlantısı olmasa bile site çalışsın "
            "diye geçici bellek desteği ekledim. Başvuru, sorgulama, harita "
            "şikayeti, karanlık mod, bildirim ve kayıt-giriş akışlarını "
            "baştan sona denedim. Projeyi internete yükleyip ayarlarını "
            "tamamladım. Telefona ana ekrana eklenebilir hale getirdim. "
            "Staj sorumlusuna canlı demo ile projeyi gösterdim ve "
            "10 günlük staj defterimi teslim ettim."
        ),
    },
]


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading("STAJ DEFTERİ", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("T.C. Dörtyol Belediyesi\nBilgi İşlem Müdürlüğü")
    run.font.size = Pt(14)
    run.bold = True

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub2.add_run(
        "\ne-Belediye Vatandaş Başvuru ve Yönetim Sistemi\n"
        "Web sitesi geliştirme · Veritabanı · Harita · Yönetici paneli\n"
    )
    run2.font.size = Pt(11)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Staj Süresi: 10 İş Günü (e-Belediye Projesi)\n").font.size = Pt(11)
    info.add_run("Stajyer: .............................................\n").font.size = Pt(11)
    info.add_run("Staj Sorumlusu: .............................................\n").font.size = Pt(11)

    doc.add_page_break()

    for day in DAYS:
        doc.add_heading(day["title"], level=1)

        kisa = doc.add_paragraph()
        k1 = kisa.add_run("Çalışma Konusu: ")
        k1.bold = True
        k1.font.size = Pt(10)
        k2 = kisa.add_run(day["konu"])
        k2.font.size = Pt(10)

        kisa2 = doc.add_paragraph()
        a1 = kisa2.add_run("Ana Hatlar: ")
        a1.bold = True
        a1.font.size = Pt(10)
        a2 = kisa2.add_run(day["ana_hatlar"])
        a2.font.size = Pt(10)

        doc.add_paragraph()

        img_path = IMG / day["image"]
        if img_path.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(img_path), width=Inches(5.2))
            cap = doc.add_paragraph(f"Şekil: {day['title']}")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].italic = True

        doc.add_paragraph()

        detail_run = doc.add_paragraph().add_run("Çalışma ile İlgili Açıklamalar:")
        detail_run.bold = True
        detail_run.font.size = Pt(11)

        para = doc.add_paragraph(day["text"])
        para.paragraph_format.line_spacing = 1.25
        para.paragraph_format.space_after = Pt(6)
        for run in para.runs:
            run.font.size = Pt(11)

        sig = doc.add_paragraph()
        sig.add_run("\nStajyer İmza: .........................    ").font.size = Pt(10)
        sig.add_run("Onay: .........................").font.size = Pt(10)

        doc.add_page_break()

    doc.save(OUT)
    print(f"Olusturuldu: {OUT}")


if __name__ == "__main__":
    main()
